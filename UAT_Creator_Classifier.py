import re
import os
import shutil
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 🛑 CONFIGURATION ZONE 🛑
# ==========================================
TEMPLATE_ROUTER_KEYWORD = "xxx"
LOG_SUBFOLDER = "UAT"

# 1. DCGW TAGS AND COMMANDS (NE8000)
DCGW_COMMAND_MAPPINGS = {
    "[Master cfcard dir]": "dir",
    "[Slave cfcard dir]": "dir slave#cfcard:/",
    "[Display startup]": "dis startup | no",
    "[Display license]": "dis license | no",
    "[Display memory-usage]": "dis memory-usage | no",
    "[Display cpu-usage]": "dis cpu-usage | no",
    "[Dis users]": "dis users",
    "[Dis ntp status]": "dis ntp status | no",
    "[Dis info-center]": "dis info-center | no",
    "[Dis int desc]": "dis int desc | no",
    "[Display optical-module extend information interface 100GE 0/10/0]": "display optical-module extend information interface 100GE 0/10/0",
    "[Display ospf peer brief]": "dis ospf peer br | no",
    "[Dis ip routing-table vpn-instance VRF_OM protocol ospf]": "dis ip routing-table vpn-instance VRF_OM protocol ospf | no",
    "[Display bgp vpnv4 all peer]": "display bgp vpnv4 all peer | no",
    "[Master display VRRP brief]": "display vrrp brief | no",
    "[Backup display VRRP brief]": "display vrrp brief | no",
    "[Display fan]": "dis fan | no",
    "[Display device]": "dis device | no"
}

# 2. SW TAGS AND COMMANDS (CE6800)
SW_COMMAND_MAPPINGS = {
    "[Master cfcard dir]": "dir",
    "[Display startup]": "dis startup | no",
    "[Display memory-usage]": "dis memory | no",
    "[Display cpu-usage]": "dis cpu | no",
    "[Display users]": "dis users",
    "[Display ntp status]": "dis ntp status | no",
    "[Display info-center]": "dis info-center | no",
    "[Display int desc]": "dis int desc | no",
    "[Display interface 100ge 2/0/0 transceiver verbose]": "dis interface 100ge1/0/1 transceiver verbose",
    "[Display vlan 4005]": "dis vlan 4005",
    "[Display mac-address vlan 4005]": "dis mac-address vlan 4005",
    "[Display dfs group 1 m-lag]": "dis dfs-group 1 m-lag",
    "[Display dfs-group 1 m-lag brief]": "dis dfs-group 1 m-lag brief",
    "[Display fan]": "dis fan | no"
}

# 3. 🌐 MULTI-SITE ROUTER LIST (NO IPs NEEDED) 🌐
MULTI_SITE_ROUTERS = {
    "A": [
     "AA"
    ],
    "B": [
     "BB"
    ],
    "C": [
     "CC"
    ]
}


# ==========================================

def extract_specific_command(log_text, command, router_name):
    if command == "N/A_MASTER_ROUTER":
        return f"[{router_name} is configured as 01 (Master). Backup VRRP log is N/A.]"
    if command == "N/A_BACKUP_ROUTER":
        return f"[{router_name} is configured as 02 (Backup). Master VRRP log is N/A.]"

    safe_command = re.escape(command)
    pattern = rf"{safe_command}[ \t\r]*\n?(.*?)(?=\n<|\n\[|\Z)"
    match = re.search(pattern, log_text, flags=re.DOTALL)

    if match:
        raw_output = match.group(1).strip()
        return f"<{router_name}> {command}\n{raw_output}"
    else:
        return f"[Command '{command}' not found in log]"


def replace_text_in_element(element, router_name, site_name):
    changes = False
    if TEMPLATE_ROUTER_KEYWORD in element.text:
        element.text = element.text.replace(TEMPLATE_ROUTER_KEYWORD, router_name)
        changes = True

    if "TOSB TOC" in element.text:
        element.text = element.text.replace("TOSB TOC", f"{site_name} TOC")
        changes = True

    return changes


def update_word_uat(filepath, router_name, log_text, command_mappings, site_name):
    try:
        doc = Document(filepath)
        changes_made = False

        def process_table_recursive(table):
            inner_changes = False
            for row in table.rows:
                for cell in row.cells:
                    # A. Standard Text Replacement (IP/Name)
                    if replace_text_in_element(cell, router_name, site_name):
                        inner_changes = True

                    # B. Check for Command Placeholders
                    for placeholder, command in command_mappings.items():
                        if placeholder in cell.text:
                            print(f"      [DEBUG] Injecting CLI log into: {placeholder}")
                            specific_output = extract_specific_command(log_text, command, router_name)

                            # ✨ PRESERVE BOX: Clear text without destroying cell formatting
                            # We clear each paragraph's text instead of cell.text = ""
                            for paragraph in cell.paragraphs:
                                paragraph.text = ""

                            # Inject the new log into the first paragraph
                            para = cell.paragraphs[0]
                            run = para.add_run(specific_output)
                            run.font.name = 'Consolas'
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text for CLI look
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT/

                            +
                            inner_changes = True

                    # C. Nested Table Check
                    if cell.tables:
                        for nested_table in cell.tables:
                            if process_table_recursive(nested_table):
                                inner_changes = True
            return inner_changes

        # Process Body, Headers, Footers
        for para in doc.paragraphs:
            if replace_text_in_element(para, router_name, site_name): changes_made = True
        for section in doc.sections:
            for container in [section.header, section.footer]:
                if container:
                    for para in container.paragraphs:
                        if replace_text_in_element(para, router_name, site_name): changes_made = True
                    for table in container.tables:
                        if process_table_recursive(table): changes_made = True

        for table in doc.tables:
            if process_table_recursive(table): changes_made = True

        if changes_made:
            doc.save(filepath)
            return True
        return False
    except Exception as e:
        print(f"    [!] Error updating Word doc: {e}")
        return False



def safe_read_log(filepath):
    encodings = ['utf-8', 'utf-16', 'windows-1252', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def process_multi_site_uat():
    base_dir = os.getcwd()
    outputs_base_folder = os.path.join(base_dir, "Command Outputs")

    template_dir_path = os.path.join(base_dir, "xxx")

    if not os.path.exists(template_dir_path):
        print(f"[!] CRITICAL ERROR: Could not find the template folder at '{template_dir_path}'.")
        return

    for site_name, routers_list in MULTI_SITE_ROUTERS.items():
        site_folder_path = os.path.join(base_dir, site_name)
        if not os.path.exists(site_folder_path):
            os.makedirs(site_folder_path)

        print(f"\n==================================================")
        print(f"🚀 Starting UAT Log Injection for Site: {site_name}")
        print(f"==================================================")

        for router_name in routers_list:
            upper_name = router_name.upper()

            is_dcgw = False
            is_sw = False

            if "GW" in upper_name:
                is_dcgw = True
                current_mappings = DCGW_COMMAND_MAPPINGS.copy()

                if "01" in upper_name:
                    print(f"\n[+] Processing DCGW (Master 01): {router_name}")
                    current_mappings["[Backup display VRRP brief]"] = "N/A_MASTER_ROUTER"
                elif "02" in upper_name:
                    print(f"\n[+] Processing DCGW (Backup 02): {router_name}")
                    current_mappings["[Master display VRRP brief]"] = "N/A_BACKUP_ROUTER"
                else:
                    print(f"\n[+] Processing DCGW: {router_name}")

                log_device_folder = "DCGW"
                file_keyword = "NE8000"

            elif "SW" in upper_name:
                is_sw = True
                current_mappings = SW_COMMAND_MAPPINGS.copy()
                log_device_folder = "SW"
                file_keyword = "CE6800"
                print(f"\n[+] Processing SW (CE6800): {router_name}")

            else:
                print(f"\n[-] Skipping {router_name} - Could not determine if GW or SW from name.")
                continue

            new_folder_path = os.path.join(site_folder_path, router_name)

            if not os.path.exists(new_folder_path):
                shutil.copytree(template_dir_path, new_folder_path)

                for filename in os.listdir(new_folder_path):
                    if filename.startswith("~$"): continue

                    old_file = os.path.join(new_folder_path, filename)

                    if is_dcgw and "CE6800" in filename:
                        os.remove(old_file_path)
                        continue
                    elif is_sw and "NE8000" in filename:
                        os.remove(old_file_path)
                        continue

                    new_filename = filename.replace("xxx", router_name)
                    new_file = os.path.join(new_folder_path, new_filename)

                    if old_file != new_file:
                        os.rename(old_file, new_file)
            else:
                print(f"    Existing folder found for: {router_name}. Checking for target file...")

            target_uat_file = None
            for filename in os.listdir(new_folder_path):
                # STRICTLY .docx ONLY
                if file_keyword.lower() in filename.lower() and filename.endswith(".docx") and not filename.startswith(
                        "~$"):
                    target_uat_file = os.path.join(new_folder_path, filename)
                    break

            if target_uat_file:
                # ✨ HERE IS THE FIX: Added site_name to the path! ✨
                log_dir_path = os.path.join(outputs_base_folder, LOG_SUBFOLDER, log_device_folder, site_name)
                log_file_path = None

                if os.path.exists(log_dir_path):
                    for potential_file in os.listdir(log_dir_path):
                        if router_name in potential_file:
                            log_file_path = os.path.join(log_dir_path, potential_file)
                            break

                if log_file_path and os.path.exists(log_file_path):
                    full_log_text = safe_read_log(log_file_path)

                    if update_word_uat(target_uat_file, router_name, full_log_text, current_mappings, site_name):
                        print(f"    -> SUCCESS: Injected mapped commands into template boxes.")
                    else:
                        print(f"    -> WARNING: Could not find placeholders in tables.")
                else:
                    print(
                        f"    -> [!] Missing Log File for {router_name} in folder: {LOG_SUBFOLDER}/{log_device_folder}/{site_name}")
            else:
                print(f"    -> [!] Missing Word Template containing keyword '{file_keyword}' in {router_name} folder.")

    print("\n✅ Multi-Site UAT Deployment Complete!")


if __name__ == "__main__":
    process_multi_site_uat()
