import re
import os
import shutil
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 🛑 CONFIGURATION ZONE 🛑
# ==========================================
TEMPLATE_IP_KEYWORD = "xxx.xxx.xxx.xxx"
TEMPLATE_ROUTER_KEYWORD = "X"
TEMPLATE_FILE_KEYWORD = "Vendor_Hardening_Checklist"
LOG_SUBFOLDER = "VHC"

# 🎯 THE VHC MAPPER (Placeholder in Word Box -> Xshell Command)
COMMAND_MAPPINGS = {
    "[LOG_IDLE_TIMEOUT]": "dis current-configuration configuration user-interface",
    "[LOG_SSH_STATUS]": "dis ssh server status",
    "[LOG_TACACS]": "dis current-configuration configuration hwtacacs",
    "[LOG_SYSLOG]": "dis cu | inc info",
    "[LOG_NTP]": "dis ntp status",
    "[LOG_ACL]": "dis current-configuration configuration acl-adv",
    "[LOG_SNMP_V3]": "dis current-configuration configuration snmp",
    "[LOG_SNMP_COMMUNITY]": "dis current-configuration configuration snmp"
}

ADDITIONAL_KEYWORDS = {
    "TOSB TOC": "TOSB TOC"
}

# 3. 🌐 MULTI-SITE ROUTER LIST & IP SLOT 🌐
MULTI_SITE_ROUTERS = {
    "A": {
        "B01": "aaa.aaa.aaa.aaa",
    },
    "B": {
        "B01": "bbb.bbb.bbb.bbb",
    },
    "C": {
        "C01": "ccc.ccc.ccc.ccc",
    }
}


# ==========================================

def extract_specific_command(log_text, command, router_name):
    safe_command = re.escape(command)
    pattern = rf"{safe_command}[ \t\r]*\n?(.*?)(?=\n<|\n\[|\Z)"
    match = re.search(pattern, log_text, flags=re.DOTALL)

    if match:
        raw_output = match.group(1).strip()
        return f"<{router_name}> {command}\n{raw_output}"
    else:
        return f"[Command '{command}' not found in log]"


def replace_text_in_element(element, ip_address, router_name):
    changes = False
    if TEMPLATE_IP_KEYWORD in element.text:
        element.text = element.text.replace(TEMPLATE_IP_KEYWORD, ip_address)
        changes = True
    if TEMPLATE_ROUTER_KEYWORD in element.text:
        element.text = element.text.replace(TEMPLATE_ROUTER_KEYWORD, router_name)
        changes = True
    for old_kw, new_kw in ADDITIONAL_KEYWORDS.items():
        if old_kw in element.text:
            element.text = element.text.replace(old_kw, new_kw)
            changes = True
    return changes


def update_word_vhc(filepath, router_name, ip_address, log_text):
    try:
        doc = Document(filepath)
        changes_made = False

        # --- 1. DEEP CLEAN REPLACEMENTS (Headers, Footers, and Body Paragraphs) ---
        for para in doc.paragraphs:
            if replace_text_in_element(para, ip_address, router_name): changes_made = True

        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        if replace_text_in_element(para, ip_address, router_name): changes_made = True
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if replace_text_in_element(cell, ip_address, router_name): changes_made = True

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if replace_text_in_element(para, ip_address, router_name): changes_made = True
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if replace_text_in_element(cell, ip_address, router_name): changes_made = True

        # --- 2. TABLE PLACEHOLDER INJECTION ---
        # (FIXED: Moved this out of the footer loop so it scans the main document!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:

                    # First, do standard IP/Router Name replacements in tables
                    if replace_text_in_element(cell, ip_address, router_name): changes_made = True

                    # Next, check for our command placeholders
                    for placeholder, command in COMMAND_MAPPINGS.items():
                        if placeholder in cell.text:
                            print(f"      [DEBUG] Found placeholder: {placeholder} -> Injecting log.")
                            specific_output = extract_specific_command(log_text, command, router_name)

                            # Wipe the placeholder text from the cell
                            cell.text = ""

                            # Add the log output with terminal styling
                            para = cell.paragraphs[0]
                            inserted_text = para.add_run(specific_output)

                            inserted_text.font.name = 'Consolas'
                            inserted_text.font.color.rgb = RGBColor(255, 255, 255)  # White Text

                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            changes_made = True

        if changes_made:
            doc.save(filepath)
            return True
        return False
    except Exception as e:
        print(f"    [!] Error updating Word doc: {e}")
        return False


def safe_read_log(filepath):
    encodings = ['utf-8', 'utf-16', 'windows-1252', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def process_vhc_word():
    base_dir = os.getcwd()
    template_dir_path = os.path.join(base_dir, "xxx")
    outputs_base_folder = os.path.join(base_dir, "Command Outputs")

    if not os.path.exists(template_dir_path):
        print("Error: Could not find template folder.")
        return

    # 🌐 Loop through every site!
    for site_name, routers_data in MULTI_SITE_ROUTERS.items():
        site_folder_path = os.path.join(base_dir, site_name)
        if not os.path.exists(site_folder_path):
            os.makedirs(site_folder_path)

        print(f"\n==================================================")
        print(f"🚀 Starting VHC Black Box Deployment for {site_name}")
        print(f"==================================================")

        for router_name, ip_address in routers_data.items():
            new_folder_path = os.path.join(site_folder_path, router_name)

            if not os.path.exists(new_folder_path):
                shutil.copytree(template_dir_path, new_folder_path)
                print(f"\n[+] Created new folder for: {router_name}")

                for filename in os.listdir(new_folder_path):
                    if filename.startswith("~$"): continue

                    new_filename = filename.replace("xxx", router_name)
                    old_file = os.path.join(new_folder_path, filename)
                    new_file = os.path.join(new_folder_path, new_filename)

                    if old_file != new_file:
                        os.rename(old_file, new_file)
            else:
                print(f"\nExisting folder found for: {router_name}. Checking for VHC file...")

            target_vhc_file = None
            for filename in os.listdir(new_folder_path):
                if TEMPLATE_FILE_KEYWORD.lower() in filename.lower() and filename.endswith(
                        ".docx") and not filename.startswith("~$"):
                    target_vhc_file = os.path.join(new_folder_path, filename)
                    break

            if target_vhc_file:
                # Path matching your flat VHC log structure: Command Outputs/VHC/TOSB
                log_dir_path = os.path.join(outputs_base_folder, LOG_SUBFOLDER, site_name)
                log_file_path = None

                if os.path.exists(log_dir_path):
                    for potential_file in os.listdir(log_dir_path):
                        if router_name in potential_file:
                            log_file_path = os.path.join(log_dir_path, potential_file)
                            break

                if log_file_path and os.path.exists(log_file_path):
                    full_log_text = safe_read_log(log_file_path)

                    if update_word_vhc(target_vhc_file, router_name, ip_address, full_log_text):
                        print(f"  -> SUCCESS: Injected mapped commands into template boxes.")
                    else:
                        print(f"  -> WARNING: Could not find placeholders in tables.")
                else:
                    print(f"  -> [!] Missing Log File for {router_name} in folder: {LOG_SUBFOLDER}/{site_name}")
            else:
                print(f"  -> [!] Missing VHC Word Template in {router_name} folder.")

    print("\n✅ Multi-Site VHC Deployment Complete!")


if __name__ == "__main__":
    process_vhc_word()
