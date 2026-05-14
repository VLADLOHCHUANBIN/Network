import re
import os
import codecs
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 🛑 CONFIGURATION ZONE
# ==========================================
SW_PING_DATA = {
    "A": {
        "B01": "aaa.aaa.aaa.aaa",
    },
    "B": {
        "B01": "bbb.bbb.bbb.bbb",
    },
    "C": {
        "C01": "ccc.ccc.ccc.ccc",
    }
}

LOG_BASE_PATH = os.path.join("Command Outputs", "UAT", "Ping Test")

def read_file_safely(path):
    # Solves encoding issues with SecureCRT/PuTTY logs
    for enc in ['utf-8', 'utf-16', 'utf-16-le', 'cp1252', 'latin-1']:
        try:
            with codecs.open(path, 'r', encoding=enc) as f:
                return f.read()
        except:
            continue
    return ""

def extract_ping(log_text, ip):
    lines = log_text.splitlines()
    capturing = False
    block = []
    for line in lines:
        if not capturing:
            if "ping" in line.lower() and ip in line:
                capturing = True
                block.append(line.strip())
        else:
            block.append(line)
            if "min/avg/max" in line.lower():
                break
            # Break if we hit a new hostname/prompt
            if re.match(r"^[<\[].*[>\]]", line.strip()) and len(block) > 3:
                break
    return "\n".join(block).strip()

def inject(path, content):
    try:
        doc = Document(path)
        def walk_tables(tables):
            f = False
            for t in tables:
                for r in t.rows:
                    for c in r.cells:
                        if c.tables:
                            if walk_tables(c.tables): f = True
                        if "[Ping Test from EOR]" in c.text:
                            # Clear all paragraphs in cell and inject white Consolas
                            c.text = ""
                            run = c.paragraphs[0].add_run(content)
                            run.font.name = 'Consolas'
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            f = True
            return f

        if walk_tables(doc.tables):
            doc.save(path)
            return True
        return False
    except:
        return False

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    print(f"DEBUG: Script base directory: {base}")

    for site, devices in SW_PING_DATA.items():
        print(f"\n🔍 SCANNING SITE: {site}")
        log_dir = os.path.join(base, LOG_BASE_PATH, site)

        if not os.path.exists(log_dir):
            print(f"❌ LOG FOLDER MISSING: {log_dir}")
            continue

        logs = {"EOR01": "", "EOR02": ""}
        for f in os.listdir(log_dir):
            if not (f.upper().endswith(".LOG") or f.upper().endswith(".TXT")): continue
            content = read_file_safely(os.path.join(log_dir, f))
            print(f"📄 Found Log File: {f} (Chars: {len(content)})")
            if "EOR01" in f.upper():
                logs["EOR01"] = content
            elif "EOR02" in f.upper():
                logs["EOR02"] = content

        for sw, ip in devices.items():
            # Peer check logic: EOR01 UAT needs EOR02 Log data.
            src_key = "EOR02" if "EOR01" in sw.upper() else "EOR01"
            log_data = logs.get(src_key, "")

            sw_dir = os.path.join(base, site, sw)
            if not os.path.exists(sw_dir):
                continue

            target = next((os.path.join(sw_dir, f) for f in os.listdir(sw_dir)
                           if "UAT" in f.upper() and f.endswith(".docx") and not f.startswith("~$")), None)

            if target and log_data:
                print(f"   [+] Processing {sw} ({ip})")
                res = extract_ping(log_data, ip)
                if res:
                    if inject(target, res):
                        print(f"      ✅ SUCCESS: Injected into {os.path.basename(target)}")
                    else:
                        print("      ⚠️ FAIL: Placeholder not found in doc.")
                else:
                    print(f"      ❌ FAIL: IP {ip} not found in {src_key} log.")
            else:
                print(f"      ❌ FAIL: Missing file/log for {sw}")

if __name__ == "__main__":
    main()
